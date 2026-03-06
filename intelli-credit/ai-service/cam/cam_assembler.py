import os
import json
import logging
import subprocess
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from .context_assembler import assemble_cam_context, CAMContext
from .persona_accountant import run_accountant_persona
from .persona_compliance import run_compliance_persona
from .persona_cro import run_cro_persona

logger = logging.getLogger(__name__)

def build_audit_trail(
    accountant_output: dict,
    compliance_output: dict,
    cro_output: dict,
    ctx: CAMContext
) -> List[dict]:
    """
    Collects all source_citations from all three personas and adds
    standard system-generated entries for ML scores and stress tests.
    Returns deduplicated list ready for the appendix table.
    """
    citations = []

    # Collect from all persona outputs
    for output in [accountant_output, compliance_output, cro_output]:
        citations.extend(output.get("source_citations", []))

    # Add standard ML citations
    citations.append({
        "claim": f"Financial Health score {ctx.score_financial:.1f}/40",
        "source": "LightGBM Model 1 — trained on 5000 synthetic companies calibrated to CRISIL benchmarks",
        "module": "Core ML Engine — Section 7",
        "confidence": "HIGH"
    })
    citations.append({
        "claim": f"Sector sentiment {ctx.sector_sentiment_score:.2f} ({ctx.sector_risk})",
        "source": "Claude sentiment scoring with Indian regulatory severity mappings (V3 fix)",
        "module": "LangGraph Research Agent — Section 6",
        "confidence": "HIGH"
    })

    # Add stress test citations
    for scenario, result in ctx.stress_summary.items():
        if isinstance(result, dict):
            citations.append({
                "claim": f"Stress test {scenario}: score → {result.get('stressed_score', 0):.1f}",
                "source": "LightGBM stress scenario engine — Section 7.8",
                "module": "Core ML Engine",
                "confidence": "HIGH"
            })

    # Deduplicate by claim text
    seen = set()
    deduped = []
    for c in citations:
        claim_text = c.get("claim", "")
        if claim_text not in seen:
            seen.add(claim_text)
            deduped.append(c)

    return deduped


def generate_docx(cam_data: dict, job_id: str):
    """Generates the Word document using python-docx."""
    doc = Document()
    
    # Header
    head = doc.add_heading(cam_data.get('company_name', 'Company Name'), level=0)
    for run in head.runs:
        run.bold = True
        
    doc.add_paragraph(f"CIN: {cam_data.get('company_cin', '')} | Sector: {cam_data.get('sector', '')}")
    doc.add_paragraph(f"Promoters: {', '.join(cam_data.get('promoters', []))}")
    doc.add_paragraph(f"Loan Requested: ₹{cam_data.get('loan_amount', 0.0):.2f} Cr")
    doc.add_paragraph(f"Assessment Date: {cam_data.get('date', datetime.utcnow().strftime('%Y-%m-%d'))}")
    
    doc.add_page_break()

    # SECTION 2: EXECUTIVE SUMMARY
    doc.add_heading("2. Executive Summary", level=1)
    
    cro_out = cam_data.get("cro_output", {})
    final_dec = cro_out.get("final_decision", "UNKNOWN")
    
    p = doc.add_paragraph("Final Decision: ")
    r = p.add_run(final_dec)
    r.bold = True
    if final_dec == "APPROVE":
        r.font.color.rgb = RGBColor(0, 128, 0)
    elif final_dec == "REJECT":
        r.font.color.rgb = RGBColor(255, 0, 0)
    else:
        r.font.color.rgb = RGBColor(255, 165, 0)
        
    doc.add_paragraph(f"Final Score: {cam_data.get('final_score', 0):.1f}/100")
    doc.add_paragraph(f"Probability of Default (Meta): {cam_data.get('pd_meta', 0)*100:.2f}%")
    
    if final_dec in ["APPROVE", "CONDITIONAL"]:
        doc.add_paragraph(f"Sanctioned Limit: ₹{cro_out.get('sanctioned_limit_crore', 0) or 0:.2f} Cr")
        doc.add_paragraph(f"Interest Rate: {cro_out.get('interest_rate_pct', 0) or 0}%")

    if cro_out.get("override_applied"):
        p_over = doc.add_paragraph("OVERRIDE APPLIED")
        for run in p_over.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        doc.add_paragraph(f"Reason: {cro_out.get('override_reason', '')}")

    doc.add_page_break()

    # SECTION 3: FINANCIAL ASSESSMENT
    doc.add_heading("3. Financial Assessment", level=1)
    acc_out = cam_data.get("accountant_output", {})
    doc.add_paragraph(acc_out.get("content", "Financial assessment not available."))

    # Note: Confidence badges formatting could be post-processed here if text contains [✅ HIGH CONFIDENCE]
    # In a full implementation, we'd regex search and colorize the runs.
    
    # SECTION 4: LEGAL AND GOVERNANCE
    doc.add_heading("4. Legal and Governance Assessment", level=1)
    comp_out = cam_data.get("compliance_output", {})
    doc.add_paragraph(comp_out.get("content", "Compliance assessment not available."))
    
    # SECTION 5: CHIEF RISK OFFICER RECOMMENDATION
    doc.add_heading("5. Chief Risk Officer Recommendation", level=1)
    doc.add_paragraph(cro_out.get("content", "CRO recommendation not available."))
    
    mand_cov = cro_out.get("mandatory_covenants", [])
    if mand_cov:
        doc.add_heading("Mandatory Covenants", level=2)
        for mc in mand_cov:
            doc.add_paragraph(str(mc), style='List Bullet')
            
    mon_trig = cro_out.get("monitoring_triggers", [])
    if mon_trig:
        doc.add_heading("Monitoring Triggers", level=2)
        for mt in mon_trig:
            doc.add_paragraph(str(mt), style='List Bullet')

    doc.add_page_break()

    # SECTION 6: OFFICER FIELD VISIT NOTES
    off_notes = cam_data.get("officer_notes_text")
    if off_notes:
        doc.add_heading("6. Officer Field Visit Notes", level=1)
        doc.add_paragraph(f"Notes:\n{off_notes}")
        
        if cam_data.get("injection_detected"):
            p_inj = doc.add_paragraph("PROMPT INJECTION ATTEMPT DETECTED — see audit log")
            for run in p_inj.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                
        doc.add_paragraph(f"Score Adjustments Applied: {cam_data.get('officer_notes_adj', {})}")

    doc.add_page_break()

    # SECTION 7: ML SCORE DASHBOARD (Table)
    doc.add_heading("7. ML Score Dashboard", level=1)
    table_scores = doc.add_table(rows=1, cols=2)
    table_scores.style = 'Table Grid'
    hdr_cells = table_scores.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Score'
    
    components = [
        ("Financial Health", f"{cam_data.get('score_financial', 0)}/40"),
        ("Credit Behaviour", f"{cam_data.get('score_behaviour', 0)}/30"),
        ("External Risk", f"{cam_data.get('score_external', 0)}/20"),
        ("Text Signals", f"{cam_data.get('score_text', 0)}/10"),
        ("Layer 1 (Rules)", f"{cam_data.get('layer1_score', 0)}/100"),
        ("Layer 2 (ML)", f"{cam_data.get('layer2_score', 0)}/100")
    ]
    for name, val in components:
        row_cells = table_scores.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = val

    doc.add_heading("SHAP Top Drivers", level=2)
    for k, v_list in cam_data.get('shap_by_model', {}).items():
        doc.add_paragraph(str(k), style='Heading 3')
        for sh in v_list[:3]:
            doc.add_paragraph(f"{sh.get('direction')}: {sh.get('human_label')} ({sh.get('shap_value')})", style='List Bullet')

    doc.add_page_break()

    # SECTION 8: AUDIT TRAIL APPENDIX
    doc.add_heading("8. Audit Trail Appendix — Source Citation Table", level=1)
    table_audit = doc.add_table(rows=1, cols=4)
    table_audit.style = 'Table Grid'
    hdr_cells = table_audit.rows[0].cells
    hdr_cells[0].text = 'CAM Claim'
    hdr_cells[1].text = 'Source'
    hdr_cells[2].text = 'Module'
    hdr_cells[3].text = 'Confidence'

    for citation in cam_data.get("audit_trail", []):
        row_cells = table_audit.add_row().cells
        row_cells[0].text = citation.get('claim', '')
        row_cells[1].text = citation.get('source', '')
        row_cells[2].text = citation.get('module', '')
        row_cells[3].text = citation.get('confidence', 'UNKNOWN')

    base_dir = f"/tmp/intelli-credit/{job_id}"
    os.makedirs(base_dir, exist_ok=True)
    docx_path = os.path.join(base_dir, "cam_final.docx")
    doc.save(docx_path)
    logger.info(f"Saved DOCX to {docx_path}")

    # Generate PDF
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", base_dir, docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        logger.info(f"Generated PDF for job {job_id}")
    except FileNotFoundError:
        logger.warning("LibreOffice not found, skipping PDF generation.")
    except subprocess.CalledProcessError as e:
        logger.warning(f"LibreOffice conversion failed: {e.stderr.decode()}")


async def generate_cam_pipeline(job_id: str):
    """
    Main background pipeline:
    1. Assemble Context
    2. Run P1 (Accountant)
    3. Run P2 (Compliance)
    4. Run P3 (CRO)
    5. Build Audit Trail
    6. Save Draft JSON
    7. Generate DOCX/PDF
    """
    logger.info(f"Starting CAM Generation for job {job_id}")
    
    ctx = await assemble_cam_context(job_id)
    
    # 1. Forensic Accountant
    acc_out = await run_accountant_persona(ctx)
    
    # 2. Compliance Officer
    comp_out = await run_compliance_persona(ctx)
    
    # 3. Chief Risk Officer
    cro_out = await run_cro_persona(ctx, acc_out, comp_out)
    
    # 4. Audit Trail
    audit_trail = build_audit_trail(acc_out, comp_out, cro_out, ctx)
    
    # 5. Build combined data dictionary
    cam_data = {
        "job_id": job_id,
        "company_name": ctx.company_name,
        "company_cin": ctx.company_cin,
        "sector": ctx.sector,
        "promoters": ctx.promoter_names,
        "loan_amount": ctx.loan_amount_requested,
        "date": datetime.utcnow().isoformat(),
        
        "final_score": ctx.final_score,
        "layer1_score": ctx.layer1_score,
        "layer2_score": ctx.layer2_score,
        "pd_meta": ctx.pd_meta,
        "score_financial": ctx.score_financial,
        "score_behaviour": ctx.score_behaviour,
        "score_external": ctx.score_external,
        "score_text": ctx.score_text,
        
        "accountant_output": acc_out,
        "compliance_output": comp_out,
        "cro_output": cro_out,
        
        "officer_notes_text": ctx.officer_notes_text,
        "officer_notes_adj": ctx.officer_notes_adj,
        "injection_detected": ctx.injection_detected,
        
        "shap_by_model": ctx.shap_by_model,
        "audit_trail": audit_trail
    }
    
    base_dir = f"/tmp/intelli-credit/{job_id}"
    os.makedirs(base_dir, exist_ok=True)
    
    draft_path = os.path.join(base_dir, "cam_draft.json")
    with open(draft_path, "w") as f:
        json.dump(cam_data, f, indent=2)
        
    logger.info(f"Saved CAM Draft JSON to {draft_path}")
    
    # 6. Generate DOCX/PDF
    generate_docx(cam_data, job_id)
    
    logger.info(f"CAM Generation COMPLETE for job {job_id}")
